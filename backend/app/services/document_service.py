import os
import uuid
from pathlib import Path
from typing import BinaryIO, Optional
from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException, status
import magic
import PyPDF2
import docx
from io import BytesIO

from app.models.document import Document, DocumentType
from app.models.user import User
from config.settings import settings

class DocumentService:
    def __init__(self):
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
        
        # Allowed file types
        self.allowed_extensions = {'.pdf', '.doc', '.docx', '.txt'}
        self.allowed_mime_types = {
            'application/pdf',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain'
        }
        self.max_file_size = 10 * 1024 * 1024  # 10MB

    def validate_file(self, file: UploadFile) -> None:
        """Validate uploaded file"""
        # Check file size
        if file.size and file.size > self.max_file_size:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File size too large. Maximum size is {self.max_file_size // (1024*1024)}MB"
            )
        
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in self.allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type not allowed. Allowed types: {', '.join(self.allowed_extensions)}"
            )

    def generate_filename(self, original_filename: str) -> str:
        """Generate unique filename"""
        file_ext = Path(original_filename).suffix.lower()
        unique_id = str(uuid.uuid4())
        return f"{unique_id}{file_ext}"

    async def save_file(self, file: UploadFile) -> tuple[str, int]:
        """Save uploaded file to disk"""
        filename = self.generate_filename(file.filename)
        file_path = self.upload_dir / filename
        
        content = await file.read()
        file_size = len(content)
        
        with open(file_path, "wb") as f:
            f.write(content)
        
        return str(file_path), file_size

    def extract_text_content(self, file_path: str, mime_type: str) -> Optional[str]:
        """Extract text content from file"""
        try:
            if mime_type == 'application/pdf':
                return self._extract_pdf_text(file_path)
            elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return self._extract_docx_text(file_path)
            elif mime_type == 'text/plain':
                return self._extract_txt_text(file_path)
        except Exception as e:
            print(f"Error extracting text from {file_path}: {str(e)}")
            return None

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text).strip()

    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()

    def get_mime_type(self, file_path: str) -> str:
        """Get MIME type of file"""
        try:
            return magic.from_file(file_path, mime=True)
        except:
            # Fallback based on extension
            ext = Path(file_path).suffix.lower()
            mime_map = {
                '.pdf': 'application/pdf',
                '.doc': 'application/msword',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.txt': 'text/plain'
            }
            return mime_map.get(ext, 'application/octet-stream')

    async def upload_document(
        self, 
        db: Session, 
        user: User, 
        file: UploadFile, 
        document_type: str
    ) -> Document:
        """Upload and process document"""
        # Validate file
        self.validate_file(file)
        
        # Save file
        file_path, file_size = await self.save_file(file)
        
        # Get MIME type
        mime_type = self.get_mime_type(file_path)
        
        # Extract text content
        content_text = self.extract_text_content(file_path, mime_type)
        
        # Create database record
        document = Document(
            user_id=user.id,
            document_type=document_type,
            filename=Path(file_path).name,
            original_filename=file.filename,
            file_path=file_path,
            content_text=content_text,
            file_size=file_size,
            mime_type=mime_type
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        return document

    def get_user_documents(self, db: Session, user: User) -> list[Document]:
        """Get all documents for a user"""
        return db.query(Document).filter(Document.user_id == user.id).all()

    def get_document_by_id(self, db: Session, user: User, document_id: int) -> Optional[Document]:
        """Get specific document by ID for a user"""
        return db.query(Document).filter(
            Document.id == document_id,
            Document.user_id == user.id
        ).first()

    def delete_document(self, db: Session, user: User, document_id: int) -> bool:
        """Delete document and file"""
        document = self.get_document_by_id(db, user, document_id)
        if not document:
            return False
        
        # Delete file from disk
        try:
            os.remove(document.file_path)
        except FileNotFoundError:
            pass  # File already deleted
        
        # Delete from database
        db.delete(document)
        db.commit()
        
        return True

# Global instance
document_service = DocumentService()